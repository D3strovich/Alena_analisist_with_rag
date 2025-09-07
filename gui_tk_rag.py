#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Tkinter GUI для Kids RAG (без игр)
- Чат + экспорт детального отчёта в Word (профиль, наблюдения, риски)
"""
from __future__ import annotations

import os, threading, uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None

from rag_chain import invoke_rag

AGE_BANDS = ["6-7", "8-10", "10-12", "12-14"]

class RagGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Kids RAG — Тестовый чат (психолог)")
        self.geometry("980x680")

        self.session_id = self._new_session_id()
        self.frames_log: List[Dict[str, Any]] = []
        self.index_dir = tk.StringVar(value=str(Path("./data/index").resolve()))
        self.age_band = tk.StringVar(value="8-10")
        self.temperature = tk.DoubleVar(value=0.3)

        self.base_url = tk.StringVar(value=os.getenv("OPENAI_BASE_URL", "http://127.0.0.1:1234/v1"))
        self.api_key  = tk.StringVar(value=os.getenv("OPENAI_API_KEY", "lm-studio"))
        self.model    = tk.StringVar(value=os.getenv("OPENAI_MODEL", "mistral-7b-instruct-v0.2"))

        self._build_ui()

    # ---------- UI ----------
    def _build_ui(self):
        top = ttk.Frame(self); top.pack(fill=tk.X, padx=8, pady=6)
        ttk.Label(top, text="Индекс:").grid(row=0, column=0, sticky="w")
        ttk.Entry(top, textvariable=self.index_dir, width=60).grid(row=0, column=1, sticky="we", padx=6)
        ttk.Button(top, text="…", width=3, command=self._browse_index).grid(row=0, column=2)
        ttk.Label(top, text="Возраст:").grid(row=0, column=3, padx=(18, 0))
        ttk.Combobox(top, textvariable=self.age_band, values=AGE_BANDS, width=8, state="readonly").grid(row=0, column=4)
        ttk.Label(top, text="T=").grid(row=0, column=5, padx=(18, 4))
        ttk.Spinbox(top, from_=0.0, to=1.5, increment=0.1, textvariable=self.temperature, width=6).grid(row=0, column=6)

        llm = ttk.Frame(self); llm.pack(fill=tk.X, padx=8, pady=(0, 6))
        ttk.Label(llm, text="BASE_URL:").grid(row=0, column=0, sticky="w")
        ttk.Entry(llm, textvariable=self.base_url, width=34).grid(row=0, column=1, padx=6)
        ttk.Label(llm, text="API_KEY:").grid(row=0, column=2, sticky="e")
        ttk.Entry(llm, textvariable=self.api_key, width=20, show="*").grid(row=0, column=3, padx=6)
        ttk.Label(llm, text="MODEL:").grid(row=0, column=4, sticky="e")
        ttk.Entry(llm, textvariable=self.model, width=28).grid(row=0, column=5, padx=6)
        ttk.Button(llm, text="Применить", command=self._apply_env).grid(row=0, column=6, padx=(8, 0))
        ttk.Button(llm, text="Новая сессия", command=self._new_session).grid(row=0, column=7, padx=(8, 0))

        mid = ttk.Frame(self); mid.pack(fill=tk.BOTH, expand=True, padx=8, pady=4)
        self.chat = ScrolledText(mid, wrap=tk.WORD, height=20); self.chat.pack(fill=tk.BOTH, expand=True)
        self.chat.insert(tk.END, f"Сессия: {self.session_id}\nДобро пожаловать! Напишите что-нибудь и нажмите Отправить.\n\n")
        self.chat.configure(state=tk.DISABLED)

        bottom = ttk.Frame(self); bottom.pack(fill=tk.X, padx=8, pady=6)
        self.entry = tk.Text(bottom, height=3, wrap=tk.WORD); self.entry.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        btns = ttk.Frame(bottom); btns.pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(btns, text="Отправить", command=self._on_send).pack(fill=tk.X)
        ttk.Button(btns, text="Экспорт в Word", command=self._export_word).pack(fill=tk.X, pady=(6, 0))

        self.status = ttk.Label(self, text="Готово", anchor="w"); self.status.pack(fill=tk.X, padx=8, pady=(0, 6))

    # ---------- handlers ----------
    def _browse_index(self):
        d = filedialog.askdirectory(title="Папка индекса FAISS")
        if d: self.index_dir.set(d)

    def _apply_env(self):
        os.environ["OPENAI_BASE_URL"] = self.base_url.get().strip()
        os.environ["OPENAI_API_KEY"]  = self.api_key.get().strip()
        os.environ["OPENAI_MODEL"]    = self.model.get().strip()
        messagebox.showinfo("OK", "Переменные окружения применены к текущему процессу.")

    def _new_session_id(self) -> str:
        return f"sess-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{uuid.uuid4().hex[:6]}"

    def _new_session(self):
        self.session_id = self._new_session_id()
        self.frames_log.clear()
        self._append_chat(f"\n— — — Новая сессия: {self.session_id} — — —\n")

    def _append_chat(self, text: str):
        self.chat.configure(state=tk.NORMAL)
        self.chat.insert(tk.END, text)
        self.chat.see(tk.END)
        self.chat.configure(state=tk.DISABLED)

    def _set_status(self, text: str):
        self.status.config(text=text); self.status.update_idletasks()

    def _on_send(self):
        msg = self.entry.get("1.0", tk.END).strip()
        if not msg: return
        self.entry.delete("1.0", tk.END)
        self._append_chat(f"Вы: {msg}\n")
        self._set_status("Думаю…")
        self.entry.configure(state=tk.DISABLED)

        def worker():
            try:
                reply, frames = invoke_rag(
                    question=msg,
                    age_band=self.age_band.get(),
                    controls={},
                    index_dir=self.index_dir.get(),
                    encoder_name="sentence-transformers/all-MiniLM-L6-v2",
                    base_url=os.getenv("OPENAI_BASE_URL"),
                    api_key=os.getenv("OPENAI_API_KEY"),
                    model=os.getenv("OPENAI_MODEL"),
                    temperature=float(self.temperature.get()),
                    max_tokens=360,
                )
            except Exception as e:
                self.after(0, lambda err=e: self._on_error(err)); return
            self.after(0, lambda: self._on_model_reply(reply, frames))

        threading.Thread(target=worker, daemon=True).start()

    def _on_model_reply(self, reply: str, frames: List[Dict[str, Any]]):
        shown = reply.strip() if reply and reply.strip() else "(модель прислала только JSON — см. ниже)"
        self._append_chat(f"Модель: {shown}\n\n")
        if frames:
            import json as _json
            ts = datetime.now().isoformat(timespec="seconds")
            stamped = []
            for fr in frames:
                fr2 = dict(fr); fr2["_ts"] = ts; stamped.append(fr2)
            self.frames_log.extend(stamped)
            pretty = _json.dumps(stamped, ensure_ascii=False, indent=2)
            self._append_chat(f"[JSON]\n{pretty}\n\n")
        self._set_status("Готово"); self.entry.configure(state=tk.NORMAL)

    def _on_error(self, e: Exception):
        self._append_chat(f"\n[ОШИБКА] {e}\n\n")
        self._set_status("Ошибка"); self.entry.configure(state=tk.NORMAL)

    # ---------- Экспорт в Word (без игр) ----------
    def _export_word(self):
        if Document is None:
            messagebox.showerror("Нет python-docx", "Установите пакет: pip install python-docx"); return
        if not self.frames_log:
            if not messagebox.askyesno("Пусто", "Фреймов нет. Сохранить пустой отчёт?"): return
        path = filedialog.asksaveasfilename(
            title="Сохранить отчёт",
            defaultextension=".docx",
            filetypes=[("Word", "*.docx"), ("Все файлы", "*.*")],
            initialfile=f"report_{self.session_id}.docx"
        )
        if not path: return
        try:
            self._build_docx(Path(path)); messagebox.showinfo("OK", f"Отчёт сохранён: {path}")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def _build_docx(self, path: Path):
        from collections import defaultdict
        Document  # ensure import

        def _extend_set(acc, vals):
            if not vals: return acc
            if isinstance(vals, (list, tuple, set)):
                acc.update([str(x) for x in vals if x])
            else:
                acc.add(str(vals))
            return acc

        # агрегаты
        moods = []                # [(ts, mood)]
        interests = set()
        school_likes = set()
        school_diff = set()
        habits = {"sleep_hours": None, "screen_time": None}
        stressors = set()
        risk_flags = set()
        events_by_theme = defaultdict(list)

        for fr in self.frames_log:
            t = str(fr.get("type", "")).lower()
            ts = fr.get("_ts")

            if t == "profile_update":
                if fr.get("mood_now"): moods.append((ts, str(fr["mood_now"])))
                _extend_set(interests, fr.get("interests"))
                school = fr.get("school") or {}
                _extend_set(school_likes, school.get("likes"))
                _extend_set(school_diff,  school.get("difficulties"))
                _hab = fr.get("habits") or {}
                habits["sleep_hours"] = _hab.get("sleep_hours") or habits["sleep_hours"]
                habits["screen_time"] = _hab.get("screen_time") or habits["screen_time"]
                _extend_set(stressors, fr.get("stressors"))
                _extend_set(risk_flags, fr.get("risk_flags"))

            elif t == "report_event":
                theme = fr.get("theme") or "other"
                summary = fr.get("summary") or ""
                events_by_theme[str(theme)].append(summary)
                if str(theme) == "risk":
                    _extend_set(risk_flags, ["reported_risk"])

        # DOCX
        doc = Document()
        h = doc.add_heading("Отчёт по сессии Kids RAG (психолог)", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph(f"Сессия: {self.session_id}")
        doc.add_paragraph(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()

        # Краткое резюме
        doc.add_heading("Краткое резюме", level=2)
        last_mood = moods[-1][1] if moods else None
        bullets = []
        if last_mood: bullets.append(f"Текущее настроение: {last_mood}.")
        if interests: bullets.append("Интересы: " + ", ".join(sorted(interests)) + ".")
        if school_diff: bullets.append("Трудности в школе: " + ", ".join(sorted(school_diff)) + ".")
        if stressors: bullets.append("Стрессоры: " + ", ".join(sorted(stressors)) + ".")
        if risk_flags:
            bullets.append("⚠️ Обнаружены рисковые признаки: " + ", ".join(sorted(risk_flags)) + ". "
                           "Порекомендуйте обратиться к взрослому/специалисту.")
        if not bullets:
            bullets.append("Данных пока немного; продолжайте диалог, чтобы накопить наблюдения.")
        for b in bullets: doc.add_paragraph(b)

        # Профиль
        doc.add_heading("Профиль (агрегировано)", level=2)
        if moods:
            doc.add_paragraph("Динамика настроения:")
            for ts, m in moods: doc.add_paragraph(f"— {ts}: {m}")
        if interests:    doc.add_paragraph("Интересы: " + ", ".join(sorted(interests)))
        if school_likes: doc.add_paragraph("Нравится в школе: " + ", ".join(sorted(school_likes)))
        if school_diff:  doc.add_paragraph("Трудности в школе: " + ", ".join(sorted(school_diff)))
        if habits["sleep_hours"] or habits["screen_time"]:
            doc.add_paragraph(
                "Привычки: " +
                ("сон=" + str(habits['sleep_hours']) + "; " if habits["sleep_hours"] else "") +
                ("экранное время=" + str(habits['screen_time']) if habits["screen_time"] else "")
            )
        if stressors: doc.add_paragraph("Стрессоры: " + ", ".join(sorted(stressors)))

        # Наблюдения (по темам)
        if events_by_theme:
            doc.add_heading("Наблюдения по темам", level=2)
            for theme, items in sorted(events_by_theme.items(), key=lambda kv: (-len(kv[1]), kv[0])):
                doc.add_paragraph(f"{theme} — {len(items)} шт.")
                for s in items[:10]: doc.add_paragraph("• " + s)

        # Риск-сигналы
        if risk_flags:
            doc.add_heading("Риск-сигналы", level=2)
            for rf in sorted(risk_flags): doc.add_paragraph("• " + rf)

        # Приложение: сырые фреймы
        doc.add_heading("Приложение: технические фреймы", level=2)
        import json as _json
        doc.add_paragraph(_json.dumps(self.frames_log, ensure_ascii=False, indent=2))

        doc.save(path)


if __name__ == "__main__":
    RagGUI().mainloop()
